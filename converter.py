# -*- coding: utf-8 -*-
"""Markdown转Docx转换器"""

import os
import re
from typing import Optional
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor


class MarkdownToDocxConverter:
    """Markdown转Docx转换器"""

    def __init__(self):
        self.document = None
        self.md_file_path = None
        self.base_path = None
        self.styles_initialized = False

    def convert(self, md_file_path: str, output_path: Optional[str] = None) -> str:
        """转换Markdown文件到Docx"""
        self.md_file_path = md_file_path
        self.base_path = os.path.dirname(os.path.abspath(md_file_path))

        # 读取Markdown内容
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 创建Word文档
        self.document = Document()
        self._init_styles()

        # 解析Markdown为HTML
        md = markdown.Markdown(
            extensions=[
                'extra',
                'codehilite',
                'tables',
                'nl2br',
                'sane_lists',
            ]
        )

        # 转换为HTML
        html = md.convert(md_content)

        # 用BeautifulSoup解析HTML
        soup = BeautifulSoup(html, 'html.parser')

        # 处理body内容
        body = soup.body if soup.body else soup

        # 处理所有子元素
        for child in body.children:
            if child.name:
                self._process_element(child)

        # 确定输出路径
        if output_path is None:
            output_path = os.path.splitext(md_file_path)[0] + '.docx'

        # 保存文档
        self.document.save(output_path)
        return output_path

    def _init_styles(self):
        """初始化文档样式"""
        if self.styles_initialized:
            return

        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        self.styles_initialized = True

    def _process_element(self, element):
        """处理单个HTML元素"""
        tag = element.name

        if tag == 'h1':
            self._add_heading(element, 1)
        elif tag == 'h2':
            self._add_heading(element, 2)
        elif tag == 'h3':
            self._add_heading(element, 3)
        elif tag == 'h4':
            self._add_heading(element, 4)
        elif tag == 'h5':
            self._add_heading(element, 5)
        elif tag == 'h6':
            self._add_heading(element, 6)
        elif tag == 'p':
            self._add_paragraph(element)
        elif tag == 'ul':
            self._add_unordered_list(element)
        elif tag == 'ol':
            self._add_ordered_list(element)
        elif tag == 'blockquote':
            self._add_blockquote(element)
        elif tag == 'pre':
            self._add_code_block(element)
        elif tag == 'table':
            self._add_table(element)
        elif tag == 'hr':
            self.document.add_paragraph('_' * 20)
        elif tag == 'img':
            self._add_image(element)

    def _get_text_content(self, element) -> str:
        """获取元素的文本内容"""
        return element.get_text()

    def _add_heading(self, element, level: int):
        """添加标题"""
        text = self._get_text_content(element)
        self.document.add_heading(text, level=level)

    def _add_paragraph(self, element):
        """添加段落"""
        p = self.document.add_paragraph()
        self._process_inline_elements(p, element)

    def _process_inline_elements(self, p, element):
        """处理行内元素（加粗、斜体等）"""
        for child in element.children:
            if isinstance(child, str):
                p.add_run(child)
            elif child.name:
                text = child.get_text()
                if text:
                    run = p.add_run(text)
                    tag = child.name

                    if tag in ('strong', 'b'):
                        run.bold = True
                    elif tag in ('em', 'i'):
                        run.italic = True
                    elif tag in ('del', 's', 'strike'):
                        run.font.strike = True
                    elif tag == 'code':
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(200, 0, 0)
                    elif tag == 'a':
                        run.font.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)
                    elif tag == 'img':
                        # 图片作为段落
                        src = child.get('src', '')
                        if src:
                            self._add_image_to_paragraph(p, src)

    def _add_unordered_list(self, element):
        """添加无序列表"""
        for li in element.find_all('li', recursive=False):
            text = li.get_text()
            self.document.add_paragraph(text, style='List Bullet')

    def _add_ordered_list(self, element):
        """添加有序列表"""
        lis = element.find_all('li', recursive=False)
        for i, li in enumerate(lis):
            text = li.get_text()
            self.document.add_paragraph(f'{i+1}. {text}', style='List Number')

    def _add_blockquote(self, element):
        """添加引用块"""
        text = element.get_text()
        p = self.document.add_paragraph(text)
        p.style = 'Quote'

    def _add_code_block(self, element):
        """添加代码块"""
        text = element.get_text()
        p = self.document.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)

    def _add_table(self, element):
        """添加表格"""
        # 获取所有行
        rows = []
        for tr in element.find_all('tr'):
            row = []
            for td in tr.find_all(['td', 'th']):
                row.append(td.get_text())
            if row:
                rows.append(row)

        if rows:
            cols = len(rows[0]) if rows else 0
            table = self.document.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'

            for i, row_data in enumerate(rows):
                for j, cell_data in enumerate(row_data):
                    if j < len(table.columns):
                        cell = table.rows[i].cells[j]
                        cell.text = cell_data

    def _add_image(self, element):
        """添加图片"""
        src = element.get('src', '')
        if src:
            self._add_image_to_paragraph(self.document.add_paragraph(), src)

    def _add_image_to_paragraph(self, p, src):
        """添加图片到段落"""
        # 处理相对路径
        if not os.path.isabs(src):
            image_path = os.path.join(self.base_path, src)
        else:
            image_path = src

        # 检查文件是否存在
        if os.path.exists(image_path):
            try:
                p.add_run().add_picture(image_path, width=Inches(6))
            except Exception as e:
                print(f"添加图片失败: {e}")


def convert_file(md_path: str, output_path: str = None) -> str:
    """便捷转换函数"""
    converter = MarkdownToDocxConverter()
    return converter.convert(md_path, output_path)


if __name__ == '__main__':
    import sys
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
        output = convert_file(md_file)
        print(f"转换完成: {output}")
